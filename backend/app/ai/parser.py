import re
import os
from datetime import datetime
from typing import Dict, Tuple
import logging
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pypdf with pdfminer fallback."""
    text = ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() or ""
        if len(text.strip()) > 100:
            return text
    except Exception as e:
        logger.warning(f"pypdf failed: {e}")
    
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(file_path)
    except Exception as e:
        logger.error(f"pdfminer failed: {e}")
    
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""

def extract_text(file_path: str) -> str:
    """Route to correct extractor based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    return ""

def extract_contact_info(text: str) -> Dict[str, str]:
    """Extract email, phone, and name from resume text."""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'(\+?1?\s?)?(\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4})'
    
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    
    email = emails[0] if emails else ""
    phone = "".join(phones[0]) if phones else ""
    
    # Try to extract name from first few lines or email
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    name = ""
    SECTION_WORDS = {'experience','education','skills','summary','objective','profile',
                     'projects','certifications','employment','history','professional',
                     'work','contact','references','awards','achievements'}
    
    # Strategy 1: First line pattern
    if lines:
        first_line = lines[0].strip()
        words = first_line.split()
        if 2 <= len(words) <= 3 and not any(c.isdigit() for c in first_line) and '@' not in first_line:
            name = first_line
    
    # Strategy 2: Search top 6 lines if first line didn't work
    if not name:
        for line in lines[:6]:
            clean = line.split('|')[0].strip()
            words = clean.split()
            if 2 <= len(words) <= 4:
                if (all(w[0].isupper() for w in words if w) and
                    not any(w.isupper() and len(w) > 2 for w in words) and
                    not any(c.isdigit() for c in clean) and
                    '@' not in clean and
                    not any(w.lower() in SECTION_WORDS for w in words)):
                    name = clean
                    break
                    
    # Strategy 3: Email fallback
    if not name and email:
        username = email.split('@')[0]
        # Clean up john.doe -> John Doe
        name = " ".join([w.capitalize() for w in re.split(r'[\._-]', username)])
    
    return {"email": email, "phone": phone, "name": name or "Candidate"}

def detect_sections(text: str) -> Dict[str, str]:
    """Detect and extract resume sections."""
    section_headers = {
        "summary": ["summary", "objective", "profile", "about"],
        "experience": ["experience", "work history", "employment", "work experience", "professional experience"],
        "education": ["education", "academic", "qualifications", "degrees"],
        "skills": ["skills", "technical skills", "competencies", "technologies", "expertise"],
        "projects": ["projects", "personal projects", "portfolio", "open source"],
        "certifications": ["certifications", "certificates", "licenses", "credentials"],
        "awards": ["awards", "achievements", "honors", "recognition"],
    }
    
    sections = {}
    lines = text.split('\n')
    current_section = "header"
    current_content = []
    
    for line in lines:
        line_lower = line.strip().lower()
        found_section = False
        for section_key, keywords in section_headers.items():
            if any(kw in line_lower for kw in keywords) and len(line.strip()) < 50:
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = section_key
                current_content = []
                found_section = True
                break
        if not found_section:
            current_content.append(line)
    
    if current_content:
        sections[current_section] = '\n'.join(current_content)
    
    return sections

def parse_years_of_experience(text: str) -> int:
    """Estimate total years of experience from resume text."""
    year_patterns = [
        r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:\w+\s*){0,4}experience',
        r'(\d+)\+?\s*years?\s+(?:of\s+)?(?:professional|relevant|industry|work|software)',
        r'(\d{4})\s*[-–]\s*(?:present|current|now|\d{4})',
    ]
    
    years_mentioned = []
    for pattern in year_patterns:
        matches = re.findall(pattern, text.lower())
        for m in matches:
            try:
                val = int(m)
                if 1 <= val <= 50:
                    years_mentioned.append(val)
                elif 1990 <= val <= datetime.now().year:
                    years_mentioned.append(datetime.now().year - val)
            except:
                pass
    
    if years_mentioned:
        return max(years_mentioned)
    return 0

def mask_pii(text: str) -> str:
    """Remove PII for fair scoring."""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'(\+?1?\s?)?(\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4})'
    
    text = re.sub(email_pattern, '[EMAIL]', text)
    text = re.sub(phone_pattern, '[PHONE]', text)
    
    return text
